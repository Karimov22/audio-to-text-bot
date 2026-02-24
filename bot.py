import os
import logging
import tempfile
import datetime
import asyncio

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters
)

import speech_recognition as sr
from pydub import AudioSegment
from docx import Document as DocxDocument
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
import edge_tts

# ===================== SOZLAMALAR =====================
BOT_TOKEN = os.environ.get("BOT_TOKEN", "BU_YERGA_TOKEN_QOYING")

LANGUAGES = {
    "🇺🇿 O'zbekcha": "uz-UZ",
    "🇷🇺 Ruscha": "ru-RU",
    "🇬🇧 Inglizcha": "en-US",
    "🇹🇷 Turkcha": "tr-TR",
}

TTS_VOICES = {
    "uz-UZ": "uz-UZ-MadinaNeural",
    "ru-RU": "ru-RU-SvetlanaNeural",
    "en-US": "en-US-GuyNeural",
    "tr-TR": "tr-TR-EmelNeural",
}

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

user_settings = {}


# ===================== YORDAMCHI FUNKSIYALAR =====================

def get_user_lang(user_id):
    return user_settings.get(user_id, {}).get("lang", "uz-UZ")


def get_user_format(user_id):
    return user_settings.get(user_id, {}).get("format", "DOCX")


def convert_to_wav(input_path):
    audio = AudioSegment.from_file(input_path)
    audio = audio.set_channels(1).set_frame_rate(16000)
    tmp = tempfile.mktemp(suffix=".wav")
    audio.export(tmp, format="wav")
    return tmp


def recognize_audio(wav_path, language="uz-UZ"):
    recognizer = sr.Recognizer()
    audio_seg = AudioSegment.from_wav(wav_path)
    duration_ms = len(audio_seg)
    chunk_ms = 30000
    texts = []

    for start in range(0, duration_ms, chunk_ms):
        end = min(start + chunk_ms, duration_ms)
        chunk = audio_seg[start:end]
        tmp_chunk = tempfile.mktemp(suffix=".wav")
        chunk.export(tmp_chunk, format="wav")
        try:
            with sr.AudioFile(tmp_chunk) as source:
                recognizer.adjust_for_ambient_noise(source, duration=0.3)
                audio_data = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio_data, language=language)
                texts.append(text)
            except sr.UnknownValueError:
                pass
        finally:
            if os.path.exists(tmp_chunk):
                os.remove(tmp_chunk)

    return " ".join(texts) if texts else None


def extract_text_from_file(file_path, ext):
    """PDF yoki DOCX dan matn ajratish"""
    if ext == ".pdf":
        doc = fitz.open(file_path)
        text = "".join([page.get_text() for page in doc])
        return text.strip()
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text.strip()
    return None


async def text_to_audio(text, lang, output_path):
    """Matnni audio ga aylantirish (edge-tts)"""
    voice = TTS_VOICES.get(lang, "uz-UZ-MadinaNeural")
    clean_text = text.replace('\n', ' ').strip()[:5000]
    communicate = edge_tts.Communicate(clean_text, voice, rate="+10%")
    await communicate.save(output_path)


def make_docx(text, source_name):
    doc = DocxDocument()
    doc.add_heading("Transkriptsiya", 0).alignment = 1
    doc.add_paragraph(f"Manba: {source_name}")
    doc.add_paragraph(f"Sana: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("─" * 50)
    doc.add_paragraph()
    doc.add_heading("Matn:", level=1)

    sentences = text.split(". ")
    buf = []
    for s in sentences:
        buf.append(s.strip())
        if len(buf) >= 3:
            doc.add_paragraph(". ".join(buf) + ".")
            buf = []
    if buf:
        doc.add_paragraph(". ".join(buf))

    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    doc.add_paragraph(f"So'zlar: {len(text.split())} | Belgilar: {len(text)}")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)
    return tmp


def make_pdf(text, source_name):
    tmp = tempfile.mktemp(suffix=".pdf")
    doc_pdf = SimpleDocTemplate(tmp, pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"], fontSize=16, alignment=1, spaceAfter=10)
    meta_style = ParagraphStyle("M", parent=styles["Normal"], fontSize=9, textColor="#888888", spaceAfter=5)
    body_style = ParagraphStyle("B", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=8)

    story = [
        Paragraph("Transkriptsiya", title_style),
        Spacer(1, 0.3*cm),
        Paragraph(f"Manba: {source_name}", meta_style),
        Paragraph(f"Sana: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style),
        Spacer(1, 0.4*cm),
        Paragraph("─" * 60, meta_style),
        Spacer(1, 0.4*cm),
        Paragraph("<b>Matn:</b>", body_style),
        Spacer(1, 0.2*cm),
    ]

    sentences = text.split(". ")
    buf = []
    for s in sentences:
        buf.append(s.strip())
        if len(buf) >= 3:
            story.append(Paragraph(". ".join(buf) + ".", body_style))
            buf = []
    if buf:
        story.append(Paragraph(". ".join(buf), body_style))

    story += [
        Spacer(1, 0.4*cm),
        Paragraph("─" * 60, meta_style),
        Paragraph(f"So'zlar: {len(text.split())} | Belgilar: {len(text)}", meta_style),
    ]
    doc_pdf.build(story)
    return tmp


# ===================== BOT HANDLERLARI =====================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    name = update.effective_user.first_name
    user_settings.setdefault(uid, {"lang": "uz-UZ", "format": "DOCX"})

    text = (
        f"👋 Salom, *{name}*!\n\n"
        "Men ikki xil vazifani bajaraman:\n\n"
        "🎙️ *Audio → Matn:*\n"
        "Ovozli xabar yoki MP3/OGG fayl yuboring → PDF/DOCX qaytaraman\n\n"
        "📄 *Fayl → Audio:*\n"
        "PDF yoki DOCX fayl yuboring → MP3 audio qaytaraman\n\n"
        "⚙️ *Sozlamalar:*\n"
        "/til — tilni tanlash\n"
        "/format — PDF yoki DOCX\n"
        "/sozlamalar — joriy sozlamalar"
    )
    await update.message.reply_text(text, parse_mode="Markdown")


async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🆘 *Yordam*\n\n"
        "🎙️ *Audio → Matn:*\n"
        "• Ovozli xabar yuboring\n"
        "• MP3, OGG, WAV, M4A fayl yuboring\n\n"
        "📄 *Fayl → Audio:*\n"
        "• PDF fayl yuboring\n"
        "• DOCX fayl yuboring\n\n"
        "📌 *Buyruqlar:*\n"
        "/til — tilni tanlash\n"
        "/format — chiqish formatini tanlash\n"
        "/sozlamalar — joriy sozlamalar",
        parse_mode="Markdown"
    )


async def settings_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    lang_name = next((k for k, v in LANGUAGES.items() if v == get_user_lang(uid)), "O'zbekcha")
    fmt = get_user_format(uid)
    await update.message.reply_text(
        f"⚙️ *Joriy sozlamalar:*\n\n"
        f"🌐 Til: {lang_name}\n"
        f"📄 Format: {fmt}",
        parse_mode="Markdown"
    )


async def choose_lang(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton(lang, callback_data=f"lang:{code}")]
        for lang, code in LANGUAGES.items()
    ]
    await update.message.reply_text(
        "🌐 Tilni tanlang:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def choose_format(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [[
        InlineKeyboardButton("📝 Word (DOCX)", callback_data="fmt:DOCX"),
        InlineKeyboardButton("📄 PDF", callback_data="fmt:PDF"),
    ]]
    await update.message.reply_text(
        "📋 Audio→Matn uchun format tanlang:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid = query.from_user.id
    await query.answer()
    user_settings.setdefault(uid, {"lang": "uz-UZ", "format": "DOCX"})

    data = query.data
    if data.startswith("lang:"):
        code = data.split(":")[1]
        user_settings[uid]["lang"] = code
        lang_name = next((k for k, v in LANGUAGES.items() if v == code), code)
        await query.edit_message_text(f"✅ Til o'rnatildi: *{lang_name}*", parse_mode="Markdown")
    elif data.startswith("fmt:"):
        fmt = data.split(":")[1]
        user_settings[uid]["format"] = fmt
        await query.edit_message_text(f"✅ Format o'rnatildi: *{fmt}*", parse_mode="Markdown")


async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Audio → Matn (PDF/DOCX)"""
    uid = update.effective_user.id
    user_settings.setdefault(uid, {"lang": "uz-UZ", "format": "DOCX"})

    audio_obj = None
    file_name = "audio"

    if update.message.voice:
        audio_obj = update.message.voice
        file_name = "voice_message"
    elif update.message.audio:
        audio_obj = update.message.audio
        file_name = update.message.audio.file_name or "audio"

    if not audio_obj:
        return

    if audio_obj.file_size and audio_obj.file_size > 20 * 1024 * 1024:
        await update.message.reply_text("❌ Fayl hajmi 20MB dan oshmasligi kerak!")
        return

    msg = await update.message.reply_text("⏳ Audio qabul qilindi, ishlanmoqda...")
    tmp_input = tmp_wav = tmp_output = None

    try:
        await msg.edit_text("📥 Fayl yuklanmoqda...")
        tg_file = await audio_obj.get_file()
        ext = os.path.splitext(file_name)[1] or ".ogg"
        tmp_input = tempfile.mktemp(suffix=ext)
        await tg_file.download_to_drive(tmp_input)

        await msg.edit_text("🔄 Audio tayyorlanmoqda...")
        tmp_wav = convert_to_wav(tmp_input)

        lang = get_user_lang(uid)
        await msg.edit_text("🎙️ Matn ajratilmoqda...")
        text = await asyncio.get_event_loop().run_in_executor(
            None, recognize_audio, tmp_wav, lang
        )

        if not text:
            await msg.edit_text(
                "❌ Matn ajratib bo'lmadi.\n"
                "• Audio sifatini tekshiring\n"
                "• /til bilan tilni o'zgartiring"
            )
            return

        fmt = get_user_format(uid)
        await msg.edit_text(f"📄 {fmt} fayl tayyorlanmoqda...")

        if fmt == "DOCX":
            tmp_output = await asyncio.get_event_loop().run_in_executor(
                None, make_docx, text, file_name
            )
            out_name = os.path.splitext(file_name)[0] + "_matn.docx"
        else:
            tmp_output = await asyncio.get_event_loop().run_in_executor(
                None, make_pdf, text, file_name
            )
            out_name = os.path.splitext(file_name)[0] + "_matn.pdf"

        await msg.delete()

        if len(text) <= 800:
            await update.message.reply_text(f"📝 *Matn:*\n\n{text}", parse_mode="Markdown")

        with open(tmp_output, "rb") as f:
            await update.message.reply_document(
                document=f,
                filename=out_name,
                caption=f"✅ *Tayyor!*\n🌐 Til: `{lang}`\n📊 So'zlar: `{len(text.split())}`\n📄 Format: `{fmt}`",
                parse_mode="Markdown"
            )

    except Exception as e:
        logger.error(f"Audio xatolik: {e}")
        await msg.edit_text(f"❌ Xatolik:\n`{str(e)}`", parse_mode="Markdown")
    finally:
        for tmp in [tmp_input, tmp_wav, tmp_output]:
            if tmp and os.path.exists(tmp):
                os.remove(tmp)


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """PDF/DOCX → Audio (MP3) yoki Audio fayl → Matn"""
    uid = update.effective_user.id
    user_settings.setdefault(uid, {"lang": "uz-UZ", "format": "DOCX"})

    doc = update.message.document
    if not doc:
        return

    file_name = doc.file_name or "fayl"
    ext = os.path.splitext(file_name)[1].lower()

    # Audio fayl bo'lsa → matnга
    if ext in [".mp3", ".wav", ".ogg", ".m4a", ".flac", ".aac"]:
        if doc.file_size and doc.file_size > 20 * 1024 * 1024:
            await update.message.reply_text("❌ Fayl hajmi 20MB dan oshmasligi kerak!")
            return

        msg = await update.message.reply_text("⏳ Audio fayl qabul qilindi...")
        tmp_input = tmp_wav = tmp_output = None

        try:
            await msg.edit_text("📥 Yuklanmoqda...")
            tg_file = await doc.get_file()
            tmp_input = tempfile.mktemp(suffix=ext)
            await tg_file.download_to_drive(tmp_input)

            await msg.edit_text("🔄 Tayyorlanmoqda...")
            tmp_wav = convert_to_wav(tmp_input)

            lang = get_user_lang(uid)
            await msg.edit_text("🎙️ Matn ajratilmoqda...")
            text = await asyncio.get_event_loop().run_in_executor(
                None, recognize_audio, tmp_wav, lang
            )

            if not text:
                await msg.edit_text("❌ Matn ajratib bo'lmadi.")
                return

            fmt = get_user_format(uid)
            if fmt == "DOCX":
                tmp_output = await asyncio.get_event_loop().run_in_executor(
                    None, make_docx, text, file_name
                )
                out_name = os.path.splitext(file_name)[0] + "_matn.docx"
            else:
                tmp_output = await asyncio.get_event_loop().run_in_executor(
                    None, make_pdf, text, file_name
                )
                out_name = os.path.splitext(file_name)[0] + "_matn.pdf"

            await msg.delete()
            if len(text) <= 800:
                await update.message.reply_text(f"📝 *Matn:*\n\n{text}", parse_mode="Markdown")

            with open(tmp_output, "rb") as f:
                await update.message.reply_document(
                    document=f, filename=out_name,
                    caption=f"✅ *Tayyor!* 📄 {fmt}",
                    parse_mode="Markdown"
                )

        except Exception as e:
            logger.error(f"Xatolik: {e}")
            await msg.edit_text(f"❌ Xatolik:\n`{str(e)}`", parse_mode="Markdown")
        finally:
            for tmp in [tmp_input, tmp_wav, tmp_output]:
                if tmp and os.path.exists(tmp):
                    os.remove(tmp)

    # PDF yoki DOCX → Audio (MP3)
    elif ext in [".pdf", ".docx"]:
        if doc.file_size and doc.file_size > 20 * 1024 * 1024:
            await update.message.reply_text("❌ Fayl hajmi 20MB dan oshmasligi kerak!")
            return

        msg = await update.message.reply_text("⏳ Fayl qabul qilindi, ishlanmoqda...")
        tmp_input = tmp_output = None

        try:
            await msg.edit_text("📥 Yuklanmoqda...")
            tg_file = await doc.get_file()
            tmp_input = tempfile.mktemp(suffix=ext)
            await tg_file.download_to_drive(tmp_input)

            await msg.edit_text("📖 Matn o'qilmoqda...")
            text = extract_text_from_file(tmp_input, ext)

            if not text or len(text.strip()) < 3:
                await msg.edit_text("❌ Faylda matn topilmadi!")
                return

            lang = get_user_lang(uid)
            await msg.edit_text("🔊 Audio tayyorlanmoqda...")
            tmp_output = tempfile.mktemp(suffix=".mp3")
            await text_to_audio(text, lang, tmp_output)

            out_name = os.path.splitext(file_name)[0] + ".mp3"
            await msg.delete()

            with open(tmp_output, "rb") as f:
                await update.message.reply_audio(
                    audio=f,
                    filename=out_name,
                    caption=f"✅ *Tayyor!*\n🌐 Til: `{lang}`\n📊 So'zlar: `{len(text.split())}`",
                    parse_mode="Markdown"
                )

        except Exception as e:
            logger.error(f"Xatolik: {e}")
            await msg.edit_text(f"❌ Xatolik:\n`{str(e)}`", parse_mode="Markdown")
        finally:
            for tmp in [tmp_input, tmp_output]:
                if tmp and os.path.exists(tmp):
                    os.remove(tmp)

    else:
        await update.message.reply_text(
            "❌ Qo'llab-quvvatlanmaydigan format!\n\n"
            "Qabul qilinadi:\n"
            "🎙️ Audio: MP3, WAV, OGG, M4A\n"
            "📄 Hujjat: PDF, DOCX"
        )


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📌 Menga quyidagilarni yuboring:\n\n"
        "🎙️ Ovozli xabar → Matn (PDF/DOCX)\n"
        "🎵 Audio fayl (MP3/OGG) → Matn\n"
        "📄 PDF/DOCX fayl → Audio (MP3)\n\n"
        "/til — tilni o'zgartirish\n"
        "/help — yordam"
    )


# ===================== MAIN =====================

def main():
    if BOT_TOKEN == "BU_YERGA_TOKEN_QOYING":
        print("❌ BOT_TOKEN o'rnatilmagan!")
        return

    app = Application.builder().token(BOT_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_cmd))
    app.add_handler(CommandHandler("til", choose_lang))
    app.add_handler(CommandHandler("format", choose_format))
    app.add_handler(CommandHandler("sozlamalar", settings_cmd))
    app.add_handler(CallbackQueryHandler(button_handler))
    app.add_handler(MessageHandler(filters.VOICE, handle_audio))
    app.add_handler(MessageHandler(filters.AUDIO, handle_audio))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    print("🤖 Bot ishga tushdi!")
    app.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()
